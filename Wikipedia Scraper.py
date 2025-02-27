from PIL import Image
import requests
import re
from io import BytesIO
import wikipedia as wiki
from docx import Document
from docx.shared import Inches

# Configurar el idioma de Wikipedia
lang = input('In what language do you want the information? ')
wiki.set_lang(lang)

# Buscar la página en Wikipedia
search = input('What do you want to search? ')
page = wiki.page(search)

folder = "destination folder"

# Crear un nuevo documento de Word
doc = Document()
doc.add_heading(page.title, level=1)

# Descargar y añadir la primera imagen de la página
try:
    response = requests.get(page.images[0])
    response.raise_for_status()
    img = Image.open(BytesIO(response.content))
    img_path = "image.jpg"
    with open(img_path, 'wb') as f:
        img.save(f)
    doc.add_picture(img_path, width=Inches(4.0))
except requests.exceptions.RequestException as e:
    print(f"Error al descargar la imagen: {e}")

summary = re.sub(r'\[.*?\]', '', page.summary)
doc.add_paragraph(summary)

# Guardar el documento
doc.save(f"{folder}{page.title}.docx")
print("Saved document")